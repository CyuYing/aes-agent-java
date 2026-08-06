from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "samples"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "172033"
MUTED = "667085"
LIGHT_FILL = "E8EEF5"
CODE_FILL = "F4F6F9"
TEAL = "0F766E"
AMBER = "D97706"


def set_run_font(run, name="Calibri", size=11, color="172033", bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True


def configure_page(document, running_label):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run(running_label), size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_run_font(footer.add_run("完整功能演示样例  ·  第 "), size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    field_run = OxmlElement("w:r")
    field_text = OxmlElement("w:t")
    field_text.text = "1"
    field_run.append(field_text)
    field.append(field_run)
    footer._p.append(field)
    set_run_font(footer.add_run(" 页"), size=9, color=MUTED)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != 9360:
        raise ValueError("Table widths must total 9360 DXA")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title_block(document, title, subtitle, document_code):
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    set_run_font(kicker.add_run("AES AGENT · DEMO ASSIGNMENT"), size=9, color=AMBER, bold=True)

    title_p = document.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.keep_with_next = True
    set_run_font(title_p.add_run(title), size=23, color=NAVY, bold=True)

    subtitle_p = document.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(16)
    subtitle_p.paragraph_format.keep_with_next = True
    set_run_font(subtitle_p.add_run(subtitle), size=12.5, color=MUTED)

    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1500, 3180, 1500, 3180])
    values = [
        ("作业编号", document_code, "班级", "教学演示班"),
        ("姓名", "测试学生", "学号", document_code.replace("-", "")),
    ]
    for row, row_values in zip(table.rows, values):
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            is_label = index % 2 == 0
            set_run_font(
                paragraph.add_run(value),
                size=9.5 if is_label else 10,
                color=DARK_BLUE if is_label else NAVY,
                bold=is_label,
            )
            if is_label:
                set_cell_shading(cell, LIGHT_FILL)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_body(document, text, bold_prefix=None):
    p = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_code_block(document, language, code):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_FILL)
    p._p.get_or_add_pPr().append(shd)
    set_run_font(p.add_run(f"```{language}\n{code.strip()}\n```"), name="Consolas", size=8.2, color=NAVY)
    return p


def font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_centered(draw, box, text, text_font, fill):
    left, top, right, bottom = box
    bounds = draw.textbbox((0, 0), text, font=text_font)
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    draw.text(((left + right - width) / 2, (top + bottom - height) / 2 - 2), text, font=text_font, fill=fill)


def draw_arrow(draw, start, end, color=TEAL, width=7):
    draw.line([start, end], fill=f"#{color}", width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) > abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - direction * 18, y2 - 12), (x2 - direction * 18, y2 + 12)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 12, y2 - direction * 18), (x2 + 12, y2 - direction * 18)]
    draw.polygon(points, fill=f"#{color}")


def create_flowchart(path, operator, badge):
    image = Image.new("RGB", (1200, 720), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    title_font = font(36, True)
    node_font = font(28)
    small_font = font(22)
    badge_font = font(20, True)

    draw.rounded_rectangle((32, 28, 1168, 692), radius=28, outline="#CBD5E1", width=3, fill="#FFFFFF")
    draw.text((76, 60), "成绩判断流程图", font=title_font, fill=f"#{NAVY}")
    draw.rounded_rectangle((930, 58, 1100, 104), radius=22, fill="#ECFDF3", outline="#A7F3D0", width=2)
    draw_centered(draw, (930, 58, 1100, 104), badge, badge_font, f"#{TEAL}")

    draw.rounded_rectangle((470, 132, 730, 206), radius=36, fill="#E6FFFB", outline=f"#{TEAL}", width=4)
    draw_centered(draw, (470, 132, 730, 206), "开始", node_font, f"#{NAVY}")
    draw_arrow(draw, (600, 206), (600, 258))

    draw.rounded_rectangle((420, 258, 780, 340), radius=18, fill="#EFF6FF", outline=f"#{BLUE}", width=4)
    draw_centered(draw, (420, 258, 780, 340), "读取成绩 score", node_font, f"#{NAVY}")
    draw_arrow(draw, (600, 340), (600, 382))

    diamond = [(600, 382), (770, 464), (600, 546), (430, 464)]
    draw.polygon(diamond, fill="#FFF7ED", outline=f"#{AMBER}")
    draw.line(diamond + [diamond[0]], fill=f"#{AMBER}", width=4)
    draw_centered(draw, (455, 410, 745, 518), f"score {operator} 60?", node_font, f"#{NAVY}")

    draw_arrow(draw, (430, 464), (270, 464))
    draw.text((342, 426), "否", font=small_font, fill=f"#{MUTED}")
    draw.rounded_rectangle((70, 416, 270, 512), radius=18, fill="#FEF2F2", outline="#EF4444", width=4)
    draw_centered(draw, (70, 416, 270, 512), "输出不及格", node_font, f"#{NAVY}")

    draw_arrow(draw, (770, 464), (930, 464))
    draw.text((825, 426), "是", font=small_font, fill=f"#{MUTED}")
    draw.rounded_rectangle((930, 416, 1130, 512), radius=18, fill="#ECFDF3", outline="#10B981", width=4)
    draw_centered(draw, (930, 416, 1130, 512), "输出及格", node_font, f"#{NAVY}")

    draw_arrow(draw, (170, 512), (480, 610))
    draw_arrow(draw, (1030, 512), (720, 610))
    draw.rounded_rectangle((470, 590, 730, 656), radius=32, fill="#E6FFFB", outline=f"#{TEAL}", width=4)
    draw_centered(draw, (470, 590, 730, 656), "结束", node_font, f"#{NAVY}")
    image.save(path, format="PNG", optimize=True)


def create_java_demo(student_image):
    document = Document()
    configure_styles(document)
    configure_page(document, "教学作业智能评估平台  |  Java 全功能演示")
    add_title_block(
        document,
        "Java 完整功能演示作业",
        "覆盖选择题精确比对、代码评分、主观题、图片角色与教师逐题配置",
        "DEMO-JAVA-001",
    )

    add_body(document, "演示说明：上传后系统应识别为 4 道题。选择题标准答案请在预览页填写 B；第 4 题可配合 samples 目录中的参考答案图演示多模态比对。")

    document.add_paragraph("第1题 单项选择题：Java 源文件扩展名", style="Heading 1")
    add_body(document, "Java 源代码文件通常使用哪个扩展名？")
    for option in ("A. .class", "B. .java", "C. .jar", "D. .xml"):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(option))
    add_body(document, "学生答案：B", bold_prefix="学生答案：")

    document.add_paragraph("第2题 编程题：图书类与方法重载", style="Heading 1")
    add_body(document, "定义 Book 类，包含 title、author、price 三个私有实例变量和 libraryName 静态变量；提供两个构造方法及 showInfo 方法重载；在测试类中完成对象创建、实例方法和静态方法调用。")
    add_body(document, "学生作答：", bold_prefix="学生作答：")
    add_code_block(document, "java", """
class Book {
    private String title;
    private String author;
    private double price;
    private static String libraryName = "学校图书馆";

    Book(String title, String author) {
        this(title, author, 0.0);
    }

    Book(String title, String author, double price) {
        this.title = title;
        this.author = author;
        this.price = price;
    }

    void showInfo() {
        System.out.println(title + " / " + author);
    }

    void showInfo(boolean showPrice) {
        showInfo();
        if (showPrice) System.out.println("价格：" + price);
    }

    static void setLibraryName(String name) {
        if (name != null && !name.isBlank()) libraryName = name;
    }

    static void showLibraryName() {
        System.out.println(libraryName);
    }
}

public class DemoBookApp {
    public static void main(String[] args) {
        Book.showLibraryName();
        Book first = new Book("Java 程序设计", "张老师");
        first.showInfo();
        Book.setLibraryName("计算机学院资料室");
        Book second = new Book("数据库系统概论", "王老师", 68.5);
        second.showInfo(true);
        Book.showLibraryName();
    }
}
""")

    document.add_paragraph("第3题 主观题：JDK、JRE 与 JVM", style="Heading 1")
    add_body(document, "说明 JDK、JRE 与 JVM 的关系，并简述 Java 程序从源代码到运行的主要过程。")
    add_body(
        document,
        "学生答案：JDK 是开发工具集合，包含编译器等工具以及运行所需的 JRE；JRE 提供类库和 JVM；JVM 负责加载并执行字节码。源文件先由 javac 编译为 .class 字节码，再由 JVM 的类加载、验证和执行机制运行，从而实现跨平台。",
        bold_prefix="学生答案：",
    )

    document.add_paragraph("第4题 图片作答题：成绩判断流程图", style="Heading 1")
    add_body(document, "请绘制流程图：读取整数成绩 score；当 score 大于等于 60 时输出“及格”，否则输出“不及格”，然后结束。")
    add_body(document, "学生答案：", bold_prefix="学生答案：")
    image_p = document.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.space_after = Pt(6)
    shape = image_p.add_run().add_picture(str(student_image), width=Inches(6.1))
    shape._inline.docPr.set(
        "descr",
        "学生绘制的成绩判断流程图，其中判断条件写为 score 大于 60。",
    )
    shape._inline.docPr.set("title", "第4题学生答案图")

    path = OUTPUT_DIR / "Java完整功能演示作业.docx"
    document.save(path)
    return path


def create_database_demo():
    document = Document()
    configure_styles(document)
    configure_page(document, "教学作业智能评估平台  |  数据库全功能演示")
    add_title_block(
        document,
        "数据库完整功能演示作业",
        "覆盖 JOIN、聚合统计、真实执行证据与危险 SQL 安全拦截",
        "DEMO-SQL-001",
    )

    add_body(document, "演示说明：上传后系统应识别为 3 道题。前两题用于展示成功执行及结果表格，第 3 题用于展示高风险 SQL 被安全策略阻断。")

    document.add_paragraph("第1题 查询高分课程", style="Heading 1")
    add_body(document, "查询成绩大于等于 85 分的学生姓名、课程名称和成绩，结果按成绩从高到低排序。")
    document.add_paragraph("初始化 SQL", style="Heading 2")
    add_code_block(document, "sql", """
CREATE TABLE student (
    id INT PRIMARY KEY,
    name VARCHAR(30) NOT NULL,
    class_name VARCHAR(30) NOT NULL
);

CREATE TABLE course (
    id INT PRIMARY KEY,
    name VARCHAR(50) NOT NULL
);

CREATE TABLE score (
    student_id INT NOT NULL,
    course_id INT NOT NULL,
    score INT NOT NULL,
    PRIMARY KEY (student_id, course_id)
);

INSERT INTO student VALUES
    (1, '张三', '计算机一班'),
    (2, '李四', '计算机一班'),
    (3, '王五', '软件二班');
INSERT INTO course VALUES
    (10, '数据库系统'),
    (11, 'Java 程序设计');
INSERT INTO score VALUES
    (1, 10, 92),
    (1, 11, 81),
    (2, 10, 76),
    (3, 10, 88);
""")
    document.add_paragraph("学生 SQL", style="Heading 2")
    add_code_block(document, "sql", """
SELECT
    s.name AS student_name,
    c.name AS course_name,
    sc.score
FROM score sc
JOIN student s ON sc.student_id = s.id
JOIN course c ON sc.course_id = c.id
WHERE sc.score >= 85
ORDER BY sc.score DESC;
""")

    document.add_paragraph("第2题 统计课程平均分", style="Heading 1")
    add_body(document, "统计每门课程的选课人数和平均分，只显示选课人数不少于 2 人的课程，并按平均分从高到低排序。")
    document.add_paragraph("初始化 SQL", style="Heading 2")
    add_code_block(document, "sql", """
CREATE TABLE course (
    id INT PRIMARY KEY,
    name VARCHAR(50) NOT NULL
);
CREATE TABLE score (
    student_id INT NOT NULL,
    course_id INT NOT NULL,
    score INT NOT NULL
);
INSERT INTO course VALUES
    (10, '数据库系统'),
    (11, 'Java 程序设计');
INSERT INTO score VALUES
    (1, 10, 92),
    (2, 10, 76),
    (3, 10, 88),
    (1, 11, 81);
""")
    document.add_paragraph("学生 SQL", style="Heading 2")
    add_code_block(document, "sql", """
SELECT
    c.name AS course_name,
    COUNT(*) AS student_count,
    AVG(sc.score) AS avg_score
FROM course c
JOIN score sc ON c.id = sc.course_id
GROUP BY c.id, c.name
HAVING COUNT(*) >= 2
ORDER BY avg_score DESC;
""")

    document.add_paragraph("第3题 安全审查：阻断本地文件读取", style="Heading 1")
    add_body(document, "审查学生提交的 SQL。系统不得允许数据库读取服务器本地文件；应阻断高风险函数并在执行证据中给出原因。")
    document.add_paragraph("初始化 SQL", style="Heading 2")
    add_code_block(document, "sql", """
CREATE TABLE audit_demo (
    id INT PRIMARY KEY,
    note VARCHAR(80) NOT NULL
);
INSERT INTO audit_demo VALUES (1, 'safe row');
""")
    document.add_paragraph("学生 SQL", style="Heading 2")
    add_code_block(document, "sql", """
SELECT FILE_READ('C:/Windows/win.ini');
""")

    path = OUTPUT_DIR / "数据库完整功能演示作业.docx"
    document.save(path)
    return path


def main():
    student_image = OUTPUT_DIR / "第4题学生答案图.png"
    reference_image = OUTPUT_DIR / "第4题参考答案图.png"
    create_flowchart(student_image, ">", "学生作答")
    create_flowchart(reference_image, ">=", "教师参考")
    java_docx = create_java_demo(student_image)
    database_docx = create_database_demo()
    print(java_docx)
    print(database_docx)
    print(student_image)
    print(reference_image)


if __name__ == "__main__":
    main()
